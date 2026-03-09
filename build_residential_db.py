#!/usr/bin/env python3
"""
IECC Residential Secretariat Database Builder
==============================================
Ingests ALL residential source data into a single SQLite database.

Sources:
  Excel:
    1. Tracking spreadsheet (88 REPC/RECP/IRCEPC proposals)
  DOCX:
    2. 216 Circulation Forms (RE/REPC/RECP proposals with votes & recommendations)

Phase model:
  - RE = Residential Energy (public input stage, CLOSED)
  - RECP = Residential Energy Code Proposal (renumbered from RE for public comment)
  - REPC = Residential Energy Public Comment (new proposals)
  - RECC = Residential Energy Consensus Committee (committee-generated)
  - IRCE = IRC Energy (public input, CLOSED)
  - IRCEPC = IRC Energy Public Comment
  - CE/CEPC = Commercial crossover proposals

⚠️ CRITICAL: The correct prefix is RECP, NOT REC. REC was an error.
   Same pattern as commercial: CEC was renamed to CECP, REC was renamed to RECP.
"""

import sqlite3
import re
import os
import hashlib
from datetime import datetime, date
from pathlib import Path

import openpyxl
from docx import Document

# ─── CONFIG ──────────────────────────────────────────────────────────────────

BASE_DIR = "/sessions/dazzling-practical-hawking/mnt/claudecomfiles"
DB_PATH = f"{BASE_DIR}/iecc_residential.db"

TRACKING_FILE = f"{BASE_DIR}/2027_RESIDENTIAL/Residential Subgroups/Residential EPLR/26-01-12 Meeting/2025 Public Comments of the 2027 Residential Numbering highlighted assigned to EPLR.xlsx"

RESIDENTIAL_DIR = f"{BASE_DIR}/2027_RESIDENTIAL"
SUBGROUPS_DIR = f"{RESIDENTIAL_DIR}/Residential Subgroups"
CONSENSUS_DIR = f"{RESIDENTIAL_DIR}/Residential Consensus Committee"

# ─── HELPERS ─────────────────────────────────────────────────────────────────

def uid(canonical_id):
    """Deterministic UID from canonical ID."""
    return hashlib.sha1(canonical_id.encode()).hexdigest()[:16]

def normalize_id(raw_id):
    """
    Normalize a residential proposal ID to canonical form.
    Handles: RE, RECP (formerly REC), REPC, RECC, IRCEPC, CE, CEPC with various formatting.
    Note: REC inputs are auto-normalized to RECP.
    Returns (canonical_id, prefix, cycle, flags[])
    """
    if not raw_id:
        return None, None, None, ["EMPTY_ID"]

    raw = str(raw_id).strip()
    # Strip leading dashes
    raw = raw.lstrip('-').strip()

    # Handle "Part I" / "Part II" / "Part III" suffixes (with or without spaces)
    part_suffix = ''
    part_match = re.search(r'\s*Part\s*(I{1,3}|IV|V)\s*$', raw, re.IGNORECASE)
    if part_match:
        part_suffix = ' Part ' + part_match.group(1).upper()
        raw = raw[:part_match.start()]

    # Handle modification suffixes: Mod, MP#, author names like "Kahre", "Schmidt"
    mod_suffix = ''
    mod_match = re.search(r'((?:Mod|MP\d*|Kahre|Schmidt|KAHRE|SCHMIDT)[A-Za-z0-9,]*)\s*$', raw)
    if mod_match:
        mod_suffix = ' ' + mod_match.group(1)
        raw = raw[:mod_match.start()].rstrip('-')

    # Clean up spaces
    raw = re.sub(r'\s+', '', raw)
    # Handle patterns like "RE-117-24" → "RE117-24"
    raw = re.sub(r'^(RE|RECP|REC|REPC|RECC|IRCEPC|IRCE|CE|CEPC|CECP)-(\d)', r'\1\2', raw, flags=re.IGNORECASE)

    flags = []

    # Standard pattern: PREFIX + NUMBER + optional -YEAR
    # Note: REC must come after RECC and RECP in the alternation to avoid partial matches
    m = re.match(r'^(IRCEPC|IRCE|REPC|RECC|RECP|REC|CEPC|CECP|CE|RE)(\d+)(?:-(\d+))?$', raw, re.IGNORECASE)
    if not m:
        flags.append(f"UNPARSEABLE_ID:{raw}")
        return raw, None, None, flags

    prefix = m.group(1).upper()
    number = m.group(2)
    year = m.group(3)

    # ⚠️ CRITICAL: REC was renamed to RECP. Normalize immediately.
    if prefix == 'REC':
        prefix = 'RECP'
        flags.append("REC_RENAMED_TO_RECP")

    # Default year based on prefix
    if not year:
        if prefix in ('REPC', 'RECP', 'RECC', 'IRCEPC', 'CEPC', 'CECP'):
            year = '25'
        else:
            year = '24'
        flags.append("MISSING_SUFFIX_FIXED")

    full_suffix = part_suffix + mod_suffix
    canonical = f"{prefix}{number}-{year}{full_suffix}".strip()
    cycle = f"20{year}" if len(year) == 2 else year

    if mod_suffix:
        flags.append("MODIFICATION_VARIANT")

    return canonical, prefix, cycle, flags

def clean_str(val):
    """Clean a cell value to string or None."""
    if val is None:
        return None
    s = str(val).strip()
    if s in ('', 'None', 'none', '-', 'N/A'):
        return None
    return s

def clean_date(val):
    """Convert various date formats to ISO string."""
    if val is None:
        return None
    if isinstance(val, datetime):
        return val.strftime('%Y-%m-%d')
    if isinstance(val, date):
        return val.isoformat()
    s = str(val).strip()
    if not s or s in ('None', '-'):
        return None
    for fmt in ('%Y-%m-%d', '%m/%d/%Y', '%m/%d/%y', '%m/%d/%Y', '%m.%d.%y'):
        try:
            return datetime.strptime(s, fmt).strftime('%Y-%m-%d')
        except ValueError:
            continue
    return s

def parse_vote_string(vote_str):
    """
    Parse various vote string formats.
    Examples:
        "8 – For Approval  1– Against Approval"
        "5 – For Disapprove; 3 – Against Disapprove; 2 – Not Voting"
        "10-0-0" (in subgroup notes)
        "? – For Disapprove; ? – Against Disapprove; ? – Not Voting"
        "Affirmative  Negative  Table" (empty template)
    Returns (vote_for, vote_against, vote_not_voting, raw_string)
    """
    if not vote_str:
        return None, None, None, None

    s = str(vote_str).strip()

    # Skip empty template lines
    if 'Affirmative' in s and 'Negative' in s:
        return None, None, None, s

    vote_for = None
    vote_against = None
    vote_not_voting = None

    # Pattern 1: "N – For Approval/Disapproval; N – Against; N – Not Voting/Abstain"
    for_match = re.search(r'(\d+)\s*[–-]\s*(?:For\s+(?:Approval|Disapprov))', s, re.IGNORECASE)
    against_match = re.search(r'(\d+)\s*[–-]\s*(?:Against\s+(?:Approval|Disapprov))', s, re.IGNORECASE)
    nv_match = re.search(r'(\d+)\s*[–-]\s*(?:Not\s+Voting|Abstain)', s, re.IGNORECASE)

    if for_match:
        vote_for = int(for_match.group(1))
    if against_match:
        vote_against = int(against_match.group(1))
    if nv_match:
        vote_not_voting = int(nv_match.group(1))

    # Pattern 2: "N-N-N" in notes (compact format)
    if vote_for is None:
        compact = re.search(r'(\d+)\s*[-–]\s*(\d+)\s*[-–]\s*(\d+)', s)
        if compact:
            vote_for = int(compact.group(1))
            vote_against = int(compact.group(2))
            vote_not_voting = int(compact.group(3))

    # Pattern 3: "N for disapproval, N against, N abstain" (in notes/reason)
    if vote_for is None:
        alt = re.search(r'(\d+)\s+for\s+\w+,?\s*(\d+)\s+against', s, re.IGNORECASE)
        if alt:
            vote_for = int(alt.group(1))
            vote_against = int(alt.group(2))
            nv2 = re.search(r'(\d+)\s+(?:abstain|not\s+voting)', s, re.IGNORECASE)
            if nv2:
                vote_not_voting = int(nv2.group(1))

    return vote_for, vote_against, vote_not_voting, s

def normalize_recommendation(raw):
    """Normalize recommendation text."""
    if not raw:
        return None, None, []

    s = str(raw).strip()
    flags = []
    lower = s.lower()

    # Remove leading "recommendation:" prefix
    lower = re.sub(r'^recommendation[:\s]+', '', lower).strip()
    s_clean = re.sub(r'^recommendation[:\s]+', '', s, flags=re.IGNORECASE).strip()

    if 'approved as submitted' in lower or 'approve as submitted' in lower or lower == 'as submitted':
        return 'Approved as Submitted', s_clean, flags
    if 'approved as modified' in lower or 'approve as modified' in lower or lower == 'as modified':
        return 'Approved as Modified', s_clean, flags
    if 'disapprov' in lower:
        if 'as submitted' in lower:
            return 'Disapproved as Submitted', s_clean, flags
        return 'Disapproved', s_clean, flags
    if 'approved' in lower or 'approve' in lower:
        return 'Approved', s_clean, flags
    if 'postponed' in lower:
        return 'Postponed', s_clean, flags
    if 'withdrawn' in lower or lower == 'wd':
        return 'Withdrawn', s_clean, flags
    if lower in ('dnp', 'do not process'):
        return 'Do Not Process', s_clean, flags
    if 'remanded' in lower:
        return 'Remanded', s_clean, flags

    if s_clean and s_clean != '-':
        flags.append(f"UNRECOGNIZED_REC:{s_clean}")
    return s_clean if s_clean and s_clean != '-' else None, s_clean, flags

# ─── SUBGROUP NAME MAPPING ──────────────────────────────────────────────────

SG_MAP = {
    'sg1': 'Consistency and Administration (SG1)',
    'sg2': 'Modeling (SG2)',
    'sg3': 'EPLR (SG3)',
    'sg4': 'Envelope (SG4)',
    'sg5': 'Existing Buildings (SG5)',
    'sg6': 'HVACR (SG6)',
    'sg7': 'Cost Effectiveness (SG7)',
    'consistency and administration': 'Consistency and Administration (SG1)',
    'admin': 'Consistency and Administration (SG1)',
    'modeling': 'Modeling (SG2)',
    'eplr': 'EPLR (SG3)',
    'electrical': 'EPLR (SG3)',
    'envelope': 'Envelope (SG4)',
    'existing': 'Existing Buildings (SG5)',
    'hvac': 'HVACR (SG6)',
    'hvacr': 'HVACR (SG6)',
    'cost': 'Cost Effectiveness (SG7)',
    'r-eplr': 'EPLR (SG3)',
    're ex': 'Existing Buildings (SG5)',
    'exs build': 'Existing Buildings (SG5)',
    'existing buildings': 'Existing Buildings (SG5)',
}

def normalize_subgroup(raw):
    """Normalize subgroup name to canonical short form."""
    if not raw:
        return None
    s = str(raw).strip()
    key = s.lower()

    # Direct match
    if key in SG_MAP:
        return SG_MAP[key]

    # Extract SG number
    sg_num = re.search(r'sg\s*(\d)', key)
    if sg_num:
        sgkey = f'sg{sg_num.group(1)}'
        if sgkey in SG_MAP:
            return SG_MAP[sgkey]

    # Keyword match
    for keyword, canonical in SG_MAP.items():
        if keyword in key:
            return canonical

    return s

def extract_sg_from_path(filepath):
    """Extract subgroup from file path."""
    path_lower = filepath.lower()

    if 'admin' in path_lower or 'consistency' in path_lower:
        return 'Consistency and Administration (SG1)'
    if 'modeling' in path_lower:
        return 'Modeling (SG2)'
    if 'eplr' in path_lower or 'electrical' in path_lower:
        return 'EPLR (SG3)'
    if 'envelope' in path_lower:
        return 'Envelope (SG4)'
    if 'existing' in path_lower:
        return 'Existing Buildings (SG5)'
    if 'hvac' in path_lower:
        return 'HVACR (SG6)'
    if 'cost' in path_lower:
        return 'Cost Effectiveness (SG7)'

    return None

def extract_date_from_path(filepath):
    """Extract meeting date from folder path like '25-03-12 Meeting'."""
    m = re.search(r'(\d{2})-(\d{2})-(\d{2})\s*(?:Meeting|meeting)', filepath)
    if m:
        yr, mo, dy = m.groups()
        return f'20{yr}-{mo}-{dy}'
    # Also try from filename
    m = re.search(r'(\d{4})-(\d{2})-(\d{2})', filepath)
    if m:
        return f'{m.group(1)}-{m.group(2)}-{m.group(3)}'
    return None

# ─── DATABASE CREATION ───────────────────────────────────────────────────────

def create_db(conn):
    c = conn.cursor()
    c.executescript(f"""
    -- ========================================================================
    -- IECC RESIDENTIAL SECRETARIAT DATABASE
    -- Built: {datetime.now().isoformat()}
    -- Mirror of commercial schema with residential-specific data
    -- ========================================================================

    DROP TABLE IF EXISTS data_quality_flags;
    DROP TABLE IF EXISTS consensus_actions;
    DROP TABLE IF EXISTS subgroup_actions;
    DROP TABLE IF EXISTS subgroup_movements;
    DROP TABLE IF EXISTS meetings;
    DROP TABLE IF EXISTS proposals;
    DROP VIEW IF EXISTS v_current_status;
    DROP VIEW IF EXISTS v_ready_for_consensus;
    DROP VIEW IF EXISTS v_full_disposition;
    DROP VIEW IF EXISTS v_data_quality_review;
    DROP VIEW IF EXISTS v_multi_action_proposals;

    CREATE TABLE proposals (
        proposal_uid    TEXT PRIMARY KEY,
        canonical_id    TEXT NOT NULL UNIQUE,
        original_id     TEXT,
        cdpaccess_id    INTEGER,
        cdpaccess_url   TEXT,
        cycle           TEXT NOT NULL DEFAULT '2025',
        prefix          TEXT,
        phase           TEXT DEFAULT 'PUBLIC_COMMENT',
        code_section    TEXT,
        proponent       TEXT,
        proponent_email TEXT,
        initial_subgroup TEXT,
        current_subgroup TEXT,
        withdrawn       INTEGER DEFAULT 0,
        withdrawn_date  TEXT,
        withdrawn_reason TEXT,
        source_file     TEXT,
        created_at      TEXT DEFAULT (datetime('now'))
    );

    CREATE TABLE subgroup_movements (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        proposal_uid    TEXT NOT NULL REFERENCES proposals(proposal_uid),
        sequence        INTEGER NOT NULL,
        move_from       TEXT,
        move_to         TEXT,
        move_date       TEXT,
        source_file     TEXT,
        UNIQUE(proposal_uid, sequence)
    );

    CREATE TABLE subgroup_actions (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        proposal_uid    TEXT NOT NULL REFERENCES proposals(proposal_uid),
        subgroup        TEXT NOT NULL,
        action_date     TEXT,
        recommendation  TEXT,
        recommendation_raw TEXT,
        vote_for        INTEGER,
        vote_against    INTEGER,
        vote_not_voting INTEGER,
        vote_raw        TEXT,
        reason          TEXT,
        modification_text TEXT,
        notes           TEXT,
        source_file     TEXT,
        UNIQUE(proposal_uid, subgroup, action_date)
    );

    CREATE TABLE consensus_actions (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        proposal_uid    TEXT NOT NULL REFERENCES proposals(proposal_uid),
        sequence        INTEGER NOT NULL DEFAULT 1,
        action_date     TEXT,
        recommendation  TEXT,
        recommendation_raw TEXT,
        vote_for        INTEGER,
        vote_against    INTEGER,
        vote_not_voting INTEGER,
        vote_raw        TEXT,
        reason          TEXT,
        modification_text TEXT,
        notes           TEXT,
        moved_by        TEXT,
        seconded_by     TEXT,
        is_final        INTEGER DEFAULT 0,
        source          TEXT,
        source_file     TEXT,
        UNIQUE(proposal_uid, sequence)
    );

    CREATE TABLE meetings (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        meeting_date    TEXT NOT NULL,
        meeting_time    TEXT,
        body            TEXT DEFAULT 'Residential Consensus Committee',
        phase           TEXT,
        status          TEXT DEFAULT 'SCHEDULED',
        tentative       INTEGER DEFAULT 0,
        action_count    INTEGER DEFAULT 0,
        notes           TEXT,
        source          TEXT,
        UNIQUE(meeting_date, body)
    );

    CREATE TABLE data_quality_flags (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        proposal_uid    TEXT,
        canonical_id    TEXT,
        table_name      TEXT,
        flag_type       TEXT NOT NULL,
        raw_value       TEXT,
        resolved_value  TEXT,
        needs_review    INTEGER DEFAULT 0,
        created_at      TEXT DEFAULT (datetime('now'))
    );

    -- Indexes
    CREATE INDEX idx_proposals_canonical ON proposals(canonical_id);
    CREATE INDEX idx_proposals_prefix ON proposals(prefix);
    CREATE INDEX idx_proposals_subgroup ON proposals(current_subgroup);
    CREATE INDEX idx_sa_proposal ON subgroup_actions(proposal_uid);
    CREATE INDEX idx_sa_subgroup ON subgroup_actions(subgroup);
    CREATE INDEX idx_ca_proposal ON consensus_actions(proposal_uid);
    CREATE INDEX idx_ca_final ON consensus_actions(is_final);
    CREATE INDEX idx_meetings_date ON meetings(meeting_date);
    CREATE INDEX idx_dqf_proposal ON data_quality_flags(proposal_uid);
    CREATE INDEX idx_dqf_review ON data_quality_flags(needs_review);
    """)
    conn.commit()
    print("[INIT] Database tables created")

# ─── INGEST: TRACKING SPREADSHEET ────────────────────────────────────────────

def ingest_tracking_spreadsheet(conn):
    """Load 88 residential proposals from tracking spreadsheet."""
    print("\n[1] Ingesting Residential Tracking Spreadsheet...")

    wb = openpyxl.load_workbook(TRACKING_FILE, read_only=True, data_only=True)
    ws = wb["IECC-Residential Tracking"]

    count = 0
    prefix_counts = {}
    sg_counts = {}

    for row in ws.iter_rows(min_row=3, values_only=True):
        raw_id = clean_str(row[0])
        if not raw_id:
            break

        canonical, prefix, cycle, flags = normalize_id(raw_id)
        if not canonical:
            continue

        p_uid = uid(canonical)

        code_section = clean_str(row[1])
        cdp_id = int(row[2]) if row[2] else None
        proponent = clean_str(row[3])
        email = clean_str(row[4])
        subgroup_raw = clean_str(row[5])
        subgroup = normalize_subgroup(subgroup_raw) if subgroup_raw else None

        # Determine phase from prefix
        if prefix in ('REPC', 'RECC', 'IRCEPC', 'CEPC'):
            phase = 'PUBLIC_COMMENT'
        elif prefix in ('RECP', 'CECP'):
            phase = 'CODE_PROPOSAL'
        elif prefix in ('RE', 'IRCE', 'CE'):
            phase = 'PUBLIC_INPUT'
        else:
            phase = 'PUBLIC_COMMENT'

        url = f"https://energy.cdpaccess.com/proposal/{cdp_id}/" if cdp_id else None

        conn.execute("""
            INSERT OR IGNORE INTO proposals
            (proposal_uid, canonical_id, original_id, cdpaccess_id, cdpaccess_url,
             cycle, prefix, phase, code_section, proponent, proponent_email,
             initial_subgroup, current_subgroup, source_file)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            p_uid, canonical, raw_id, cdp_id, url,
            cycle or '2025', prefix, phase, code_section, proponent, email,
            subgroup, subgroup, 'tracking_spreadsheet'
        ))

        for flag in flags:
            conn.execute("""
                INSERT INTO data_quality_flags
                (proposal_uid, canonical_id, table_name, flag_type, raw_value, resolved_value)
                VALUES (?, ?, 'proposals', ?, ?, ?)
            """, (p_uid, canonical, flag, raw_id, canonical))

        prefix_counts[prefix] = prefix_counts.get(prefix, 0) + 1
        if subgroup:
            sg_counts[subgroup] = sg_counts.get(subgroup, 0) + 1
        count += 1

    conn.commit()
    wb.close()

    print(f"    → {count} proposals loaded from tracking spreadsheet")
    for k, v in sorted(prefix_counts.items(), key=lambda x: str(x[0])):
        print(f"      {k}: {v}")
    print(f"    → Subgroup distribution:")
    for k, v in sorted(sg_counts.items(), key=lambda x: str(x[0])):
        print(f"      {k}: {v}")

# ─── INGEST: CIRCULATION FORMS ───────────────────────────────────────────────

def extract_from_circulation_form(docx_path):
    """
    Extract ALL proposal data from a circulation form DOCX.
    Some forms contain multiple tables (one per proposal).
    Returns list of dicts.
    """
    results = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        return results

    for table in doc.tables:
        if len(table.rows) < 8:
            continue

        data = {}
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key_text = row.cells[0].text.strip().lower()
            val_text = row.cells[1].text.strip()

            if not val_text or val_text == '-':
                # Check if key cell has the proposal ID (some forms put it there)
                if not key_text or key_text in ('proposal #', 'proposal#', ''):
                    continue

            # Proposal number - only from "Proposal #" row (NOT "Proposal Status")
            is_proposal_row = (key_text in ('proposal #', 'proposal#', 'proposal  #', ''))
            has_proposal_pattern = bool(re.search(r'(?:RE|RECP|REC|REPC|RECC|IRCEPC|CE|CEPC|CECP)\s*-?\s*\d+', val_text, re.IGNORECASE)) or bool(re.match(r'^\d+-\d+$', val_text.replace(' ', '')))

            if is_proposal_row and (val_text and has_proposal_pattern):
                pid = val_text
                # Clean up "RE 156-24" → "RE156-24"
                pid = re.sub(r'\s+', '', pid)
                # If no prefix (just "145-24"), try to get prefix from filename
                if re.match(r'^\d+-\d+$', pid):
                    fname = Path(docx_path).name
                    prefix_match = re.search(r'(IRCEPC|REPC|RECC|RECP|REC|CEPC|CECP|CE|RE)', fname, re.IGNORECASE)
                    if prefix_match:
                        pid = prefix_match.group(1).upper() + pid
                data['proposal_id_raw'] = pid

            elif 'cdp' in key_text:
                try:
                    data['cdp_id'] = int(re.search(r'\d+', val_text).group())
                except:
                    pass

            elif key_text in ('code section(s)', 'code section'):
                data['code_section'] = val_text

            elif key_text == 'proponent':
                data['proponent'] = val_text

            elif key_text == 'subgroup':
                data['subgroup_raw'] = val_text

            elif 'subgroup notes' in key_text or key_text == 'subgroup notes':
                data['subgroup_notes'] = val_text

            elif 'recommendation' in key_text and 'reason' in key_text:
                data['recommendation_reason'] = val_text

            elif key_text == 'vote' and 'Affirmative' not in val_text:
                # This is the subgroup vote row (not the consensus template)
                data['vote_raw'] = val_text

            elif 'recommendation date' in key_text:
                data['rec_date_raw'] = val_text

            elif key_text == 'committee response':
                if val_text and val_text != '-':
                    data['committee_response'] = val_text

            elif key_text == 'proposal status':
                if val_text and val_text != '-':
                    data['proposal_status'] = val_text

        if 'proposal_id_raw' in data:
            results.append(data)

    return results

def parse_recommendation_reason(text):
    """
    Parse combined recommendation/reason field from Row 9 of circulation forms.
    This field can contain:
      - Just a recommendation: "APPROVE AS MODIFIED"
      - Just a reason: "Editorial changes for clarification..."
      - Both: "Recommendation: disapprove\n5 for...\nReason: ..."
      - Both inline: "disapprove\nReason statement: ..."
    Returns (recommendation_text, reason_text)
    """
    if not text:
        return None, None

    text = text.strip()
    if text == '-':
        return None, None

    recommendation = None
    reason = None

    # Pattern 1: "Recommendation: X\n...Reason: Y"
    rec_match = re.search(r'(?:Recommendation)[:\s]+([^\n]+)', text, re.IGNORECASE)
    reason_match = re.search(r'(?:Reason(?:\s+statement)?)[:\s]+(.+)', text, re.IGNORECASE | re.DOTALL)

    if rec_match:
        recommendation = rec_match.group(1).strip()
        # Clean vote info from recommendation line
        recommendation = re.sub(r'\d+\s+for\s+\w+.*', '', recommendation, flags=re.IGNORECASE).strip()

    if reason_match:
        reason = reason_match.group(1).strip()

    # Pattern 2: First line is a recognizable recommendation keyword
    if not recommendation:
        first_line = text.split('\n')[0].strip()
        lower_first = first_line.lower()
        rec_keywords = ['approved as submitted', 'approved as modified', 'approve as submitted',
                        'approve as modified', 'disapproved', 'disapprove', 'approved', 'approve',
                        'postponed', 'withdrawn', 'do not process', 'disapprove as submitted']
        for kw in rec_keywords:
            if lower_first.startswith(kw) or lower_first == kw:
                recommendation = first_line
                # Rest is reason
                remaining = '\n'.join(text.split('\n')[1:]).strip()
                if remaining and not reason:
                    reason = remaining
                break

    # Pattern 3: No recognizable recommendation - the whole text is a reason
    if not recommendation:
        lower_text = text.lower()
        has_rec_keyword = any(kw in lower_text for kw in
                            ['approved', 'approve', 'disapproved', 'disapprove',
                             'postponed', 'withdrawn'])
        if not has_rec_keyword:
            # This is purely a reason statement, no recommendation
            reason = text
            return None, reason

    return recommendation, reason

def extract_votes_from_notes(notes_text):
    """Try to extract votes from subgroup notes text."""
    if not notes_text:
        return None, None, None

    # Look for "N-N-N" pattern
    m = re.search(r'(\d+)\s*[-–]\s*(\d+)\s*[-–]\s*(\d+)', notes_text)
    if m:
        return int(m.group(1)), int(m.group(2)), int(m.group(3))

    return None, None, None

def ingest_circulation_forms(conn):
    """Mine ALL DOCX circulation forms for subgroup actions."""
    print("\n[2] Ingesting Circulation Forms...")

    # Find ALL circulation form DOCX files anywhere in the residential tree
    circulation_files = []
    for root, dirs, files in os.walk(RESIDENTIAL_DIR):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~'):
                fname_lower = f.lower()
                if 'circulation' in fname_lower or 'circ form' in fname_lower:
                    circulation_files.append(os.path.join(root, f))

    print(f"    → Found {len(circulation_files)} circulation form files")

    proposals_created = 0
    actions_created = 0
    actions_skipped = 0
    errors = 0

    seen_actions = set()  # Track (proposal_uid, subgroup) to avoid duplicates

    for filepath in sorted(circulation_files):
        results = extract_from_circulation_form(filepath)

        for data in results:
            raw_pid = data.get('proposal_id_raw', '')
            canonical, prefix, cycle, id_flags = normalize_id(raw_pid)
            if not canonical:
                errors += 1
                continue

            p_uid = uid(canonical)

            # Ensure proposal exists in DB
            existing = conn.execute("SELECT 1 FROM proposals WHERE proposal_uid = ?", (p_uid,)).fetchone()
            if not existing:
                # Create proposal from circulation form data
                if prefix in ('RE', 'IRCE', 'CE'):
                    phase = 'PUBLIC_INPUT'
                elif prefix in ('RECP', 'CECP'):
                    phase = 'CODE_PROPOSAL'
                else:
                    phase = 'PUBLIC_COMMENT'
                subgroup = normalize_subgroup(data.get('subgroup_raw')) or extract_sg_from_path(filepath)

                conn.execute("""
                    INSERT OR IGNORE INTO proposals
                    (proposal_uid, canonical_id, original_id, cdpaccess_id, cycle, prefix, phase,
                     code_section, proponent, initial_subgroup, current_subgroup, source_file)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    p_uid, canonical, raw_pid,
                    data.get('cdp_id'),
                    cycle or '2024', prefix, phase,
                    data.get('code_section'),
                    data.get('proponent'),
                    subgroup, subgroup,
                    filepath
                ))
                proposals_created += 1

                for flag in id_flags:
                    conn.execute("""
                        INSERT INTO data_quality_flags
                        (proposal_uid, canonical_id, table_name, flag_type, raw_value, resolved_value)
                        VALUES (?, ?, 'proposals', ?, ?, ?)
                    """, (p_uid, canonical, flag, raw_pid, canonical))

            # Extract subgroup action data
            subgroup = normalize_subgroup(data.get('subgroup_raw')) or extract_sg_from_path(filepath)
            action_date = None

            # Try recommendation date first, then path date
            if data.get('rec_date_raw'):
                action_date = clean_date(data['rec_date_raw'])
            if not action_date:
                action_date = extract_date_from_path(filepath)

            # Parse recommendation and reason from the combined field
            rec_text = None
            reason = None

            if data.get('recommendation_reason'):
                rec_text, reason = parse_recommendation_reason(data['recommendation_reason'])

            recommendation, rec_raw, rec_flags = normalize_recommendation(rec_text)

            # If no recommendation from rec/reason field, try subgroup notes
            if not recommendation and data.get('subgroup_notes'):
                notes = data['subgroup_notes'].lower()
                if 'motion to approve as modified' in notes or 'approved as modified' in notes:
                    recommendation = 'Approved as Modified'
                elif 'motion to approve as submitted' in notes or 'approved as submitted' in notes:
                    recommendation = 'Approved as Submitted'
                elif 'motion to approve' in notes:
                    recommendation = 'Approved'
                elif 'motion to disapprove' in notes or 'disapprove' in notes:
                    recommendation = 'Disapproved'

            # Parse votes
            vote_for, vote_against, vote_nv, vote_raw = parse_vote_string(data.get('vote_raw'))

            # If no votes from vote field, try subgroup notes
            if vote_for is None and data.get('subgroup_notes'):
                vf, va, vnv = extract_votes_from_notes(data['subgroup_notes'])
                if vf is not None:
                    vote_for, vote_against, vote_nv = vf, va, vnv

            # Also try to extract votes from recommendation_reason if embedded
            if vote_for is None and data.get('recommendation_reason'):
                vf, va, vnv = extract_votes_from_notes(data['recommendation_reason'])
                if vf is not None:
                    vote_for, vote_against, vote_nv = vf, va, vnv

            # Skip if no meaningful data
            if not recommendation and vote_for is None:
                actions_skipped += 1
                continue

            # Avoid duplicate actions
            action_key = (p_uid, subgroup or 'UNKNOWN', action_date or 'NODATE')
            if action_key in seen_actions:
                continue
            seen_actions.add(action_key)

            # Modification text from subgroup notes
            mod_text = None
            if data.get('subgroup_notes') and 'modified' in data.get('subgroup_notes', '').lower():
                mod_text = data['subgroup_notes']

            try:
                conn.execute("""
                    INSERT INTO subgroup_actions
                    (proposal_uid, subgroup, action_date, recommendation, recommendation_raw,
                     vote_for, vote_against, vote_not_voting, vote_raw,
                     reason, modification_text, notes, source_file)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    p_uid, subgroup or 'UNKNOWN', action_date,
                    recommendation, rec_raw,
                    vote_for, vote_against, vote_nv, vote_raw or data.get('vote_raw'),
                    reason,
                    mod_text,
                    data.get('subgroup_notes'),
                    filepath
                ))
                actions_created += 1
            except sqlite3.IntegrityError:
                actions_skipped += 1

            for flag in rec_flags:
                conn.execute("""
                    INSERT INTO data_quality_flags
                    (proposal_uid, canonical_id, table_name, flag_type, raw_value, needs_review)
                    VALUES (?, ?, 'subgroup_actions', ?, ?, 1)
                """, (p_uid, canonical, flag, rec_text))

    conn.commit()

    print(f"    → {proposals_created} new proposals created from circulation forms")
    print(f"    → {actions_created} subgroup actions created")
    print(f"    → {actions_skipped} skipped (no data or duplicate)")
    print(f"    → {errors} parse errors")

# ─── INGEST: CONSENSUS COMMITTEE MINUTES ──────────────────────────────────────

def extract_consensus_from_docx(docx_path):
    """
    Extract consensus committee actions from meeting minutes DOCX.
    Returns list of action dicts.
    """
    actions = []
    try:
        doc = Document(docx_path)
    except:
        return actions

    full_text = '\n'.join(p.text for p in doc.paragraphs)

    # Pattern: look for proposal IDs followed by action context
    proposal_pattern = r'((?:RE|REPC|RECC|RECP|REC|IRCEPC|CEPC|CECP|CE)\s*-?\s*\d+\s*-?\s*\d*)'

    lines = full_text.split('\n')
    for i, line in enumerate(lines):
        matches = re.finditer(proposal_pattern, line, re.IGNORECASE)
        for match in matches:
            raw_pid = match.group(1)
            canonical, prefix, cycle, flags = normalize_id(raw_pid)
            if not canonical:
                continue

            # Get surrounding context (3 lines before and after)
            context_start = max(0, i - 3)
            context_end = min(len(lines), i + 5)
            context = ' '.join(lines[context_start:context_end]).lower()

            action_data = {'proposal_id': canonical, 'raw_id': raw_pid}

            # Look for action keywords in context
            if 'approved as modified' in context or 'approve as modified' in context:
                action_data['recommendation'] = 'Approved as Modified'
            elif 'approved as submitted' in context or 'approve as submitted' in context:
                action_data['recommendation'] = 'Approved as Submitted'
            elif 'disapproved' in context or 'disapprove' in context:
                action_data['recommendation'] = 'Disapproved'
            elif 'approved' in context:
                action_data['recommendation'] = 'Approved'
            elif 'postponed' in context:
                action_data['recommendation'] = 'Postponed'
            elif 'withdrawn' in context:
                action_data['recommendation'] = 'Withdrawn'

            # Look for vote counts in context
            vote_match = re.search(r'(\d+)\s*[-–]\s*(\d+)\s*[-–]\s*(\d+)', context)
            if vote_match:
                action_data['vote_for'] = int(vote_match.group(1))
                action_data['vote_against'] = int(vote_match.group(2))
                action_data['vote_not_voting'] = int(vote_match.group(3))

            # Look for mover/seconder
            motion_match = re.search(r'motion\s+(?:by|from)\s+(\w+(?:\s+\w+)?)', context)
            if motion_match:
                action_data['moved_by'] = motion_match.group(1).title()

            second_match = re.search(r'second(?:ed)?\s+(?:by)?\s+(\w+(?:\s+\w+)?)', context)
            if second_match:
                action_data['seconded_by'] = second_match.group(1).title()

            if 'recommendation' in action_data:
                actions.append(action_data)

    return actions

def ingest_consensus_minutes(conn):
    """Mine consensus committee meeting minutes for consensus actions."""
    print("\n[3] Ingesting Consensus Committee Minutes...")

    # Find all minutes-type DOCX files
    minutes_files = []
    for root, dirs, files in os.walk(CONSENSUS_DIR):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~'):
                fname_lower = f.lower()
                if 'minute' in fname_lower or 'agenda' in fname_lower:
                    minutes_files.append(os.path.join(root, f))

    # Also check subgroup folders for minutes
    for root, dirs, files in os.walk(SUBGROUPS_DIR):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~'):
                fname_lower = f.lower()
                if 'minute' in fname_lower and 'approved' in fname_lower:
                    minutes_files.append(os.path.join(root, f))

    print(f"    → Found {len(minutes_files)} minutes/agenda files")

    action_count = 0
    seen_consensus = set()

    for filepath in sorted(minutes_files):
        meeting_date = extract_date_from_path(filepath)
        actions = extract_consensus_from_docx(filepath)

        # Record meeting
        if meeting_date:
            body = 'Residential Consensus Committee'
            if 'subgroup' in filepath.lower() or 'sg' in filepath.lower():
                body = extract_sg_from_path(filepath) or 'Subgroup'

            try:
                conn.execute("""
                    INSERT OR IGNORE INTO meetings
                    (meeting_date, body, status, source)
                    VALUES (?, ?, 'COMPLETED', ?)
                """, (meeting_date, body, filepath))
            except:
                pass

        for action_data in actions:
            canonical = action_data['proposal_id']
            p_uid = uid(canonical)

            # Check if proposal exists
            existing = conn.execute("SELECT 1 FROM proposals WHERE proposal_uid = ?", (p_uid,)).fetchone()
            if not existing:
                continue

            # Avoid duplicates
            consensus_key = (p_uid, action_data.get('recommendation'), meeting_date)
            if consensus_key in seen_consensus:
                continue
            seen_consensus.add(consensus_key)

            # Get next sequence
            max_seq = conn.execute(
                "SELECT COALESCE(MAX(sequence), 0) FROM consensus_actions WHERE proposal_uid = ?",
                (p_uid,)
            ).fetchone()[0]

            try:
                conn.execute("""
                    INSERT INTO consensus_actions
                    (proposal_uid, sequence, action_date, recommendation, recommendation_raw,
                     vote_for, vote_against, vote_not_voting,
                     moved_by, seconded_by, is_final, source, source_file)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1, 'MINUTES', ?)
                """, (
                    p_uid, max_seq + 1, meeting_date,
                    action_data.get('recommendation'),
                    action_data.get('recommendation'),
                    action_data.get('vote_for'),
                    action_data.get('vote_against'),
                    action_data.get('vote_not_voting'),
                    action_data.get('moved_by'),
                    action_data.get('seconded_by'),
                    filepath
                ))
                action_count += 1
            except sqlite3.IntegrityError:
                pass

    conn.commit()
    print(f"    → {action_count} consensus actions extracted")
    print(f"    → {len(minutes_files)} files processed")

# ─── POST-PROCESSING: DATA QUALITY FLAGS ──────────────────────────────────────

def generate_data_quality_flags(conn):
    """Generate data quality flags for missing or suspicious data."""
    print("\n[4] Generating Data Quality Flags...")

    c = conn.cursor()
    flag_count = 0

    # Proposals with no subgroup action
    for row in c.execute("""
        SELECT p.proposal_uid, p.canonical_id
        FROM proposals p
        WHERE p.withdrawn = 0
          AND p.proposal_uid NOT IN (SELECT proposal_uid FROM subgroup_actions)
          AND p.proposal_uid NOT IN (SELECT proposal_uid FROM consensus_actions)
    """):
        conn.execute("""
            INSERT INTO data_quality_flags
            (proposal_uid, canonical_id, table_name, flag_type, needs_review)
            VALUES (?, ?, 'proposals', 'NO_ACTION_RECORDED', 1)
        """, (row[0], row[1]))
        flag_count += 1

    # Subgroup actions with no votes
    for row in c.execute("""
        SELECT sa.proposal_uid, p.canonical_id
        FROM subgroup_actions sa
        JOIN proposals p ON sa.proposal_uid = p.proposal_uid
        WHERE sa.vote_for IS NULL AND sa.vote_against IS NULL
    """):
        conn.execute("""
            INSERT INTO data_quality_flags
            (proposal_uid, canonical_id, table_name, flag_type, needs_review)
            VALUES (?, ?, 'subgroup_actions', 'MISSING_VOTE_COUNT', 0)
        """, (row[0], row[1]))
        flag_count += 1

    # Subgroup actions with no recommendation
    for row in c.execute("""
        SELECT sa.proposal_uid, p.canonical_id
        FROM subgroup_actions sa
        JOIN proposals p ON sa.proposal_uid = p.proposal_uid
        WHERE sa.recommendation IS NULL
    """):
        conn.execute("""
            INSERT INTO data_quality_flags
            (proposal_uid, canonical_id, table_name, flag_type, needs_review)
            VALUES (?, ?, 'subgroup_actions', 'MISSING_RECOMMENDATION', 1)
        """, (row[0], row[1]))
        flag_count += 1

    # Proposals with no proponent email (from tracking spreadsheet)
    for row in c.execute("""
        SELECT p.proposal_uid, p.canonical_id
        FROM proposals p
        WHERE p.proponent_email IS NULL AND p.source_file = 'tracking_spreadsheet'
    """):
        conn.execute("""
            INSERT INTO data_quality_flags
            (proposal_uid, canonical_id, table_name, flag_type, needs_review)
            VALUES (?, ?, 'proposals', 'MISSING_PROPONENT_EMAIL', 0)
        """, (row[0], row[1]))
        flag_count += 1

    conn.commit()
    print(f"    → {flag_count} data quality flags generated")

# ─── CREATE VIEWS ────────────────────────────────────────────────────────────

def create_views(conn):
    """Create analysis views mirroring commercial schema."""
    print("\n[5] Creating database views...")

    c = conn.cursor()
    c.executescript("""
    -- Current status of each proposal (matches commercial v_current_status)
    CREATE VIEW IF NOT EXISTS v_current_status AS
    SELECT
        p.proposal_uid,
        p.canonical_id,
        p.prefix,
        p.phase,
        p.proponent,
        p.proponent_email,
        p.current_subgroup,
        p.code_section,
        p.cdpaccess_id,
        p.withdrawn,
        sa.recommendation AS sg_recommendation,
        sa.action_date AS sg_date,
        sa.vote_for AS sg_vote_for,
        sa.vote_against AS sg_vote_against,
        sa.vote_not_voting AS sg_vote_nv,
        ca.recommendation AS consensus_recommendation,
        ca.action_date AS consensus_date,
        ca.vote_for AS ca_vote_for,
        ca.vote_against AS ca_vote_against,
        ca.vote_not_voting AS ca_vote_nv,
        CASE
            WHEN p.withdrawn = 1 THEN 'WITHDRAWN'
            WHEN ca.recommendation IS NOT NULL THEN 'DECIDED'
            WHEN p.phase = 'PUBLIC_INPUT' THEN 'PHASE_CLOSED'
            WHEN sa.recommendation IS NOT NULL THEN 'PENDING_CONSENSUS'
            ELSE 'PENDING'
        END AS computed_status
    FROM proposals p
    LEFT JOIN (
        SELECT *, ROW_NUMBER() OVER (PARTITION BY proposal_uid ORDER BY action_date DESC, id DESC) AS rn
        FROM subgroup_actions
    ) sa ON p.proposal_uid = sa.proposal_uid AND sa.rn = 1
    LEFT JOIN (
        SELECT *, ROW_NUMBER() OVER (PARTITION BY proposal_uid ORDER BY sequence DESC, id DESC) AS rn
        FROM consensus_actions WHERE is_final = 1
    ) ca ON p.proposal_uid = ca.proposal_uid AND ca.rn = 1;

    -- Proposals ready for consensus hearing
    CREATE VIEW IF NOT EXISTS v_ready_for_consensus AS
    SELECT
        p.canonical_id,
        p.proponent,
        p.current_subgroup,
        sa.recommendation AS sg_recommendation,
        sa.vote_for || '-' || sa.vote_against || '-' || sa.vote_not_voting AS sg_vote,
        sa.action_date AS sg_date
    FROM proposals p
    JOIN subgroup_actions sa ON p.proposal_uid = sa.proposal_uid
    WHERE p.withdrawn = 0
      AND p.proposal_uid NOT IN (
          SELECT proposal_uid FROM consensus_actions WHERE is_final = 1
      )
    ORDER BY p.current_subgroup, p.canonical_id;

    -- Full disposition report
    CREATE VIEW IF NOT EXISTS v_full_disposition AS
    SELECT
        p.canonical_id,
        p.prefix,
        p.phase,
        p.proponent,
        p.proponent_email,
        p.current_subgroup,
        p.code_section,
        p.withdrawn,
        sa.subgroup AS sg_body,
        sa.recommendation AS sg_recommendation,
        sa.vote_for || '-' || sa.vote_against || '-' || sa.vote_not_voting AS sg_vote,
        sa.action_date AS sg_date,
        sa.reason AS sg_reason,
        ca.recommendation AS consensus_recommendation,
        ca.vote_for || '-' || ca.vote_against || '-' || ca.vote_not_voting AS consensus_vote,
        ca.action_date AS consensus_date,
        ca.reason AS consensus_reason,
        CASE
            WHEN p.withdrawn = 1 THEN 'WITHDRAWN'
            WHEN ca.recommendation IS NOT NULL THEN ca.recommendation
            WHEN sa.recommendation IS NOT NULL THEN 'PENDING_CONSENSUS: ' || sa.recommendation
            ELSE 'NO_ACTION'
        END AS final_status
    FROM proposals p
    LEFT JOIN subgroup_actions sa ON p.proposal_uid = sa.proposal_uid
    LEFT JOIN (
        SELECT * FROM consensus_actions WHERE is_final = 1
    ) ca ON p.proposal_uid = ca.proposal_uid
    ORDER BY p.canonical_id;

    -- Data quality review
    CREATE VIEW IF NOT EXISTS v_data_quality_review AS
    SELECT
        dqf.id,
        dqf.canonical_id,
        dqf.table_name,
        dqf.flag_type,
        dqf.raw_value,
        dqf.resolved_value,
        dqf.needs_review,
        dqf.created_at
    FROM data_quality_flags dqf
    ORDER BY dqf.needs_review DESC, dqf.created_at DESC;

    -- Proposals with multiple consensus actions (procedural chains)
    CREATE VIEW IF NOT EXISTS v_multi_action_proposals AS
    SELECT
        p.canonical_id,
        p.proponent,
        ca.sequence,
        ca.action_date,
        ca.recommendation,
        ca.vote_for || '-' || ca.vote_against || '-' || ca.vote_not_voting AS vote,
        ca.is_final
    FROM consensus_actions ca
    JOIN proposals p ON ca.proposal_uid = p.proposal_uid
    WHERE p.proposal_uid IN (
        SELECT proposal_uid FROM consensus_actions GROUP BY proposal_uid HAVING COUNT(*) > 1
    )
    ORDER BY p.canonical_id, ca.sequence;
    """)

    conn.commit()
    print("    → Views created: v_current_status, v_ready_for_consensus, v_full_disposition, v_data_quality_review, v_multi_action_proposals")

# ─── VERIFY ──────────────────────────────────────────────────────────────────

def verify_db(conn):
    """Run verification queries and print summary."""
    print("\n" + "=" * 80)
    print("DATABASE VERIFICATION REPORT")
    print("=" * 80)

    c = conn.cursor()

    # Overall counts
    total_proposals = c.execute("SELECT COUNT(*) FROM proposals").fetchone()[0]
    total_sa = c.execute("SELECT COUNT(*) FROM subgroup_actions").fetchone()[0]
    total_ca = c.execute("SELECT COUNT(*) FROM consensus_actions").fetchone()[0]
    total_meetings = c.execute("SELECT COUNT(*) FROM meetings").fetchone()[0]
    total_flags = c.execute("SELECT COUNT(*) FROM data_quality_flags").fetchone()[0]

    print(f"\n  Total proposals:       {total_proposals}")
    print(f"  Subgroup actions:      {total_sa}")
    print(f"  Consensus actions:     {total_ca}")
    print(f"  Meetings:              {total_meetings}")
    print(f"  Data quality flags:    {total_flags}")

    # By prefix
    print("\n  Proposals by prefix:")
    for row in c.execute("SELECT prefix, COUNT(*) FROM proposals GROUP BY prefix ORDER BY prefix"):
        print(f"    {row[0] or 'NULL':8s}: {row[1]:3d}")

    # By phase
    print("\n  Proposals by phase:")
    for row in c.execute("SELECT phase, COUNT(*) FROM proposals GROUP BY phase ORDER BY phase"):
        print(f"    {row[0] or 'NULL':20s}: {row[1]:3d}")

    # By subgroup
    print("\n  Proposals by subgroup:")
    for row in c.execute("SELECT current_subgroup, COUNT(*) FROM proposals WHERE current_subgroup IS NOT NULL GROUP BY current_subgroup ORDER BY current_subgroup"):
        print(f"    {row[0]:40s}: {row[1]:3d}")

    # Status summary
    print("\n  Status summary:")
    for row in c.execute("SELECT computed_status, COUNT(*) FROM v_current_status GROUP BY computed_status ORDER BY computed_status"):
        print(f"    {row[0]:25s}: {row[1]:3d}")

    # Subgroup action recommendations
    if total_sa > 0:
        print("\n  Subgroup action recommendations:")
        for row in c.execute("SELECT recommendation, COUNT(*) FROM subgroup_actions WHERE recommendation IS NOT NULL GROUP BY recommendation ORDER BY COUNT(*) DESC"):
            print(f"    {row[0]:30s}: {row[1]:3d}")

    # Vote coverage
    sa_with_votes = c.execute("SELECT COUNT(*) FROM subgroup_actions WHERE vote_for IS NOT NULL").fetchone()[0]
    sa_with_reason = c.execute("SELECT COUNT(*) FROM subgroup_actions WHERE reason IS NOT NULL AND reason != ''").fetchone()[0]
    print(f"\n  SG action vote coverage:   {sa_with_votes}/{total_sa} ({sa_with_votes*100//max(total_sa,1)}%)")
    print(f"  SG action reason coverage: {sa_with_reason}/{total_sa} ({sa_with_reason*100//max(total_sa,1)}%)")

    # Withdrawn
    withdrawn = c.execute("SELECT COUNT(*) FROM proposals WHERE withdrawn = 1").fetchone()[0]
    print(f"\n  Withdrawn proposals: {withdrawn}")

    # Email coverage
    with_email = c.execute("SELECT COUNT(*) FROM proposals WHERE proponent_email IS NOT NULL").fetchone()[0]
    print(f"  Proponent email coverage: {with_email}/{total_proposals} ({with_email*100//max(total_proposals,1)}%)")

    # Data quality flag summary
    print("\n  Data quality flags by type:")
    for row in c.execute("SELECT flag_type, COUNT(*) FROM data_quality_flags GROUP BY flag_type ORDER BY COUNT(*) DESC LIMIT 10"):
        print(f"    {row[0]:40s}: {row[1]:3d}")

    print("\n" + "=" * 80)

# ─── MAIN ────────────────────────────────────────────────────────────────────

def main():
    print("=" * 80)
    print("IECC RESIDENTIAL DATABASE BUILDER")
    print(f"Started: {datetime.now().isoformat()}")
    print("=" * 80)

    if not os.path.exists(TRACKING_FILE):
        print(f"ERROR: Tracking file not found: {TRACKING_FILE}")
        return

    if os.path.exists(DB_PATH):
        os.remove(DB_PATH)
        print(f"Removed existing database")

    conn = sqlite3.connect(DB_PATH)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")

    try:
        create_db(conn)
        ingest_tracking_spreadsheet(conn)
        ingest_circulation_forms(conn)
        ingest_consensus_minutes(conn)
        generate_data_quality_flags(conn)
        create_views(conn)
        verify_db(conn)

        size_kb = os.path.getsize(DB_PATH) / 1024
        print(f"\n✓ Database created: {DB_PATH}")
        print(f"✓ Size: {size_kb:.1f} KB")
        print(f"✓ Completed: {datetime.now().isoformat()}")

    finally:
        conn.close()

if __name__ == "__main__":
    main()
